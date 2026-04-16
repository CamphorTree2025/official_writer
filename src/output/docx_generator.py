"""Word文档生成模块"""
import re
from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from ..config import get_config


class DocxGenerator:
    """Word文档生成器，保持模板格式替换变量"""

    def __init__(self):
        self.config = get_config()
        self.pattern = re.compile(self.config.variable_pattern)

    def generate(
        self,
        template_path: Path,
        filled_content: Dict[str, str],
        output_path: Path,
    ):
        """
        生成Word文档

        Args:
            template_path: 模板文件路径
            filled_content: {变量名: 填充内容}
            output_path: 输出文件路径
        """
        doc = Document(template_path)
        self._replace_in_document(doc, filled_content)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _replace_in_document(self, doc: Document, filled_content: Dict[str, str]):
        """替换文档中的变量"""
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, filled_content)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_paragraph(cell, filled_content, is_cell=True)

    def _replace_in_paragraph(self, para, filled_content: Dict[str, str], is_cell: bool = False):
        """替换段落中的变量"""
        full_text = para.text

        for var_name, content in filled_content.items():
            placeholder = f"{{{{{var_name}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, content)

        if full_text != para.text:
            if is_cell:
                para.text = full_text
            else:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = full_text
                else:
                    para.add_run(full_text)
