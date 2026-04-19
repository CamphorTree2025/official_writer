"""模板解析模块"""
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document

from ..config import get_config


class TemplateParser:
    """解析Word模板文件，提取占位符"""

    def __init__(self):
        self.config = get_config()
        self.pattern = re.compile(self.config.variable_pattern)

    def parse(self, template_path: Path) -> Tuple[str, List[str]]:
        """
        解析模板文件

        Args:
            template_path: 模板文件路径(.docx)

        Returns:
            (纯文本内容, 变量名列表)  变量按模板中出现顺序排列
        """
        doc = Document(template_path)
        full_text = []
        seen = set()
        ordered_variables = []

        for paragraph in doc.paragraphs:
            text = paragraph.text
            full_text.append(text)
            variables = self.extract_variables(text)
            for v in variables:
                if v not in seen:
                    seen.add(v)
                    ordered_variables.append(v)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    full_text.append(text)
                    variables = self.extract_variables(text)
                    for v in variables:
                        if v not in seen:
                            seen.add(v)
                            ordered_variables.append(v)

        return "\n".join(full_text), ordered_variables

    def extract_variables(self, text: str) -> List[str]:
        """
        从文本中提取变量名

        Args:
            text: 包含{{变量名}}的文本

        Returns:
            变量名列表
        """
        matches = self.pattern.findall(text)
        return [m.strip() for m in matches if m.strip()]

    def get_variable_positions(self, template_path: Path) -> List[dict]:
        """
        获取变量在文档中的位置信息（用于后续替换）

        Args:
            template_path: 模板文件路径

        Returns:
            包含变量位置信息的列表
        """
        doc = Document(template_path)
        positions = []

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            for match in self.pattern.finditer(text):
                positions.append({
                    "type": "paragraph",
                    "para_idx": para_idx,
                    "start": match.start(),
                    "end": match.end(),
                    "var_name": match.group(1).strip(),
                    "full_match": match.group(0),
                    "run": None,
                })

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    for match in self.pattern.finditer(text):
                        positions.append({
                            "type": "table",
                            "table_idx": table_idx,
                            "row_idx": row_idx,
                            "cell_idx": cell_idx,
                            "start": match.start(),
                            "end": match.end(),
                            "var_name": match.group(1).strip(),
                            "full_match": match.group(0),
                        })

        return positions

    def get_preview(self, template_path: Path) -> dict:
        """
        获取模板预览信息

        Args:
            template_path: 模板文件路径

        Returns:
            包含文本内容和变量信息的字典，变量按模板中出现顺序排列
        """
        doc = Document(template_path)
        paragraphs = []
        tables = []
        seen = set()
        ordered_variables = []

        # 解析段落
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                # 标记变量位置
                preview_text = text
                for match in self.pattern.finditer(text):
                    v = match.group(1).strip()
                    if v not in seen:
                        seen.add(v)
                        ordered_variables.append(v)
                paragraphs.append({
                    "text": text,
                    "has_variable": bool(self.extract_variables(text))
                })

        # 解析表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    text = cell.text
                    if text.strip():
                        for match in self.pattern.finditer(text):
                            v = match.group(1).strip()
                            if v not in seen:
                                seen.add(v)
                                ordered_variables.append(v)
                        row_data.append({
                            "text": text,
                            "has_variable": bool(self.extract_variables(text))
                        })
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "variables": ordered_variables,
            "variable_count": len(ordered_variables)
        }
