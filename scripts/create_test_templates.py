"""生成测试公文模板"""
import sys
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, name_cn="仿宋", name_en="Times New Roman", size=16):
    """设置 run 的字体"""
    run.font.size = Pt(size)
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def add_title(doc, text, font_cn="方正小标宋简体", font_en="Times New Roman", size=22):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size)
    run.bold = True
    return p


def add_body(doc, text, font_cn="仿宋", size=16, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, font_cn, "Times New Roman", size)
    return p


def add_variable_paragraph(doc, text, font_cn="仿宋", size=16, indent=True):
    """添加包含变量的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, font_cn, "Times New Roman", size)
    return p


def create_notice_template():
    """创建通知模板"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    add_title(doc, "关于召开{{会议名称}}的通知")
    
    add_variable_paragraph(doc, "{{发文机关}}", font_cn="仿宋", size=16, indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_variable_paragraph(doc, "{{发文字号}}", font_cn="仿宋", size=16, indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, "各部门、各单位：")

    add_variable_paragraph(doc, "{{正文内容}}")

    add_body(doc, "请各部门按照要求做好相关准备工作，确保会议顺利进行。")

    add_body(doc, "")
    add_body(doc, "")

    add_variable_paragraph(doc, "{{落款单位}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_variable_paragraph(doc, "{{日期}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def create_minutes_template():
    """创建会议纪要模板"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    add_title(doc, "{{会议名称}}会议纪要")

    # 会议信息表格
    table = doc.add_table(rows=5, cols=2, style='Table Grid')
    table.autofit = True

    info = [
        ("会议时间", "{{会议时间}}"),
        ("会议地点", "{{会议地点}}"),
        ("主 持 人", "{{主持人}}"),
        ("参会人员", "{{参会人员}}"),
        ("记 录 人", "{{记录人}}"),
    ]
    for i, (label, value) in enumerate(info):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        for cell in [cell0, cell1]:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, "仿宋", "Times New Roman", 14)

    doc.add_paragraph()  # 空行

    add_body(doc, "会议议题：")
    add_variable_paragraph(doc, "{{会议议题}}")

    add_body(doc, "会议内容：")
    add_variable_paragraph(doc, "{{会议内容}}")

    add_body(doc, "会议决议：")
    add_variable_paragraph(doc, "{{会议决议}}")

    doc.add_paragraph()
    add_variable_paragraph(doc, "{{落款单位}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_variable_paragraph(doc, "{{日期}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def create_request_template():
    """创建请示模板"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    add_title(doc, "关于{{请示事项}}的请示")

    add_variable_paragraph(doc, "{{主送机关}}：")

    add_variable_paragraph(doc, "{{请示缘由}}")

    add_variable_paragraph(doc, "{{具体事项}}")

    add_body(doc, "妥否，请批示。")

    doc.add_paragraph()
    add_variable_paragraph(doc, "{{落款单位}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_variable_paragraph(doc, "{{日期}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def create_report_template():
    """创建工作报告模板"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    add_title(doc, "关于{{报告事项}}的报告")

    add_variable_paragraph(doc, "{{主送机关}}：")

    add_body(doc, "一、工作背景")
    add_variable_paragraph(doc, "{{工作背景}}")

    add_body(doc, "二、工作开展情况")
    add_variable_paragraph(doc, "{{工作开展情况}}")

    add_body(doc, "三、取得成效")
    add_variable_paragraph(doc, "{{取得成效}}")

    add_body(doc, "四、存在问题")
    add_variable_paragraph(doc, "{{存在问题}}")

    add_body(doc, "五、下一步工作计划")
    add_variable_paragraph(doc, "{{下一步计划}}")

    doc.add_paragraph()
    add_variable_paragraph(doc, "{{落款单位}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_variable_paragraph(doc, "{{日期}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def create_letter_template():
    """创建函模板"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    add_title(doc, "关于{{函件事项}}的函")

    add_variable_paragraph(doc, "{{主送机关}}：")

    add_variable_paragraph(doc, "{{函件缘由}}")

    add_variable_paragraph(doc, "{{具体事项}}")

    add_body(doc, "请予支持为盼。")

    doc.add_paragraph()
    add_variable_paragraph(doc, "{{落款单位}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_variable_paragraph(doc, "{{日期}}", indent=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


if __name__ == "__main__":
    template_dir = project_root / "templates"
    template_dir.mkdir(exist_ok=True)

    templates = [
        ("通知模板.docx", create_notice_template),
        ("会议纪要模板.docx", create_minutes_template),
        ("请示模板.docx", create_request_template),
        ("工作报告模板.docx", create_report_template),
        ("函模板.docx", create_letter_template),
    ]

    for name, creator in templates:
        path = template_dir / name
        if not path.exists():  # 不覆盖已有模板
            doc = creator()
            doc.save(str(path))
            print(f"✅ 已创建模板: {name}")
        else:
            print(f"⏭ 模板已存在，跳过: {name}")

    print(f"\n📁 模板目录: {template_dir}")
